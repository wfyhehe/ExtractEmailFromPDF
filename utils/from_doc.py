#!/usr/bin/env python
# -*- coding: utf-8 -*-

import re
import docx


def extract_email_from_doc(doc_name):
    """
    doc_name='xxx.doc(x)'
    return a list of emails
    :param doc_name:
    :return list['xxx@xxx.xxx'...]:
    """
    block_text = ''
    doc = docx.Document(doc_name)
    docText = '\n\n'.join([
        paragraph.text.encode('utf-8') for paragraph in doc.paragraphs
    ])
    paras = doc.paragraphs

    for p in paras:
        block_text += p.text + '\n'

    print block_text
    print docText
    email_regex = u'([a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+)'
    return re.findall(email_regex, block_text)


def extract_cellphone_from_doc(doc_name):
    """
    doc_name='xxx.doc(x)'
    return a list of cellphone numbers
    :param doc_name:
    :return list['xxxxxxxxxx'...]:
    """
    block_text = ''
    doc = docx.Document(doc_name)
    paras = doc.paragraphs
    for p in paras:
        block_text += p.text

    phone_regex = u'\D([1][3,4,5,7,8][0-9]{9})\D'
    return re.findall(phone_regex, block_text)


if __name__ == '__main__':
    print extract_email_from_doc('email.docx')
    print extract_cellphone_from_doc('phone.docx')
